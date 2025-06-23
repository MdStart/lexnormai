from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from sqlalchemy.orm import Session
from typing import List, Optional
import aiofiles
import os
import io
from ...core.database import get_db
from ...core.config import settings
from ...models.models import CourseContent
from ...schemas.schemas import (
    CourseContentCreate,
    CourseContentUpdate,
    CourseContentResponse,
    FileUploadResponse
)
from ...services.gemini_service import gemini_service

# Import PDF and document processing libraries
try:
    from PyPDF2 import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

router = APIRouter()


async def extract_text_from_pdf(file_content: bytes) -> str:
    """Extract text from PDF file."""
    if not PDF_AVAILABLE:
        raise HTTPException(
            status_code=500,
            detail="PDF processing not available. PyPDF2 not installed."
        )
    
    try:
        # Create a PDF reader from bytes
        pdf_file = io.BytesIO(file_content)
        pdf_reader = PdfReader(pdf_file)
        
        # Extract text from all pages
        text_content = ""
        for page_num, page in enumerate(pdf_reader.pages):
            try:
                page_text = page.extract_text()
                if page_text:
                    text_content += f"\n--- Page {page_num + 1} ---\n"
                    text_content += page_text + "\n"
            except Exception as e:
                text_content += f"\n--- Page {page_num + 1} (Error reading page: {str(e)}) ---\n"
        
        if not text_content.strip():
            raise HTTPException(
                status_code=422,
                detail="No text content could be extracted from the PDF file."
            )
        
        return text_content.strip()
        
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Error extracting text from PDF: {str(e)}"
        )


async def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from DOCX file."""
    if not DOCX_AVAILABLE:
        raise HTTPException(
            status_code=500,
            detail="DOCX processing not available. python-docx not installed."
        )
    
    try:
        # Create a document from bytes
        docx_file = io.BytesIO(file_content)
        doc = Document(docx_file)
        
        # Extract text from all paragraphs
        text_content = ""
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content += " | ".join(row_text) + "\n"
        
        if not text_content.strip():
            raise HTTPException(
                status_code=422,
                detail="No text content could be extracted from the DOCX file."
            )
        
        return text_content.strip()
        
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Error extracting text from DOCX: {str(e)}"
        )


async def save_upload_file(upload_file: UploadFile) -> str:
    """Save uploaded file and return the content as text."""
    
    # Read file content
    file_content = await upload_file.read()
    
    # Extract text based on file type
    filename_lower = upload_file.filename.lower()
    
    if filename_lower.endswith('.txt') or filename_lower.endswith('.md'):
        # Plain text or markdown files
        try:
            text_content = file_content.decode('utf-8')
        except UnicodeDecodeError:
            try:
                text_content = file_content.decode('latin-1')
            except UnicodeDecodeError:
                raise HTTPException(
                    status_code=422,
                    detail="Unable to decode text file. Please ensure it's in UTF-8 or Latin-1 encoding."
                )
    
    elif filename_lower.endswith('.pdf'):
        # PDF files
        text_content = await extract_text_from_pdf(file_content)
    
    elif filename_lower.endswith('.docx'):
        # DOCX files
        text_content = await extract_text_from_docx(file_content)
    
    else:
        # Try to decode as text for other file types
        try:
            text_content = file_content.decode('utf-8')
        except UnicodeDecodeError:
            raise HTTPException(
                status_code=415,
                detail=f"Unsupported file type: {upload_file.filename}. Supported types: .txt, .md, .pdf, .docx"
            )
    
    return text_content


@router.post("/upload", response_model=CourseContentResponse)
async def upload_course_content(
    file: UploadFile = File(...),
    title: str = Form(...),
    db: Session = Depends(get_db)
):
    """Upload course content file and create a new course content record."""
    
    # Validate file size
    if file.size and file.size > settings.max_upload_size:
        raise HTTPException(
            status_code=413,
            detail=f"File size {file.size} exceeds maximum allowed size {settings.max_upload_size}"
        )
    
    # Validate file type
    allowed_extensions = [".txt", ".md", ".pdf", ".docx"]
    file_extension = None
    if file.filename:
        file_extension = os.path.splitext(file.filename.lower())[1]
    
    if not file_extension or file_extension not in allowed_extensions:
        raise HTTPException(
            status_code=415,
            detail=f"File type not supported. Allowed types: {', '.join(allowed_extensions)}"
        )
    
    try:
        # Extract content from file
        content = await save_upload_file(file)
        
        # Validate that we got some content
        if not content.strip():
            raise HTTPException(
                status_code=422,
                detail="The uploaded file appears to be empty or contains no extractable text."
            )
        
        # Create course content record
        course_content_data = CourseContentCreate(title=title, content=content)
        db_content = CourseContent(**course_content_data.dict())
        
        db.add(db_content)
        db.commit()
        db.refresh(db_content)
        
        return db_content
        
    except HTTPException:
        # Re-raise HTTP exceptions as-is
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error processing file: {str(e)}")


@router.post("/", response_model=CourseContentResponse)
async def create_course_content(
    content: CourseContentCreate,
    db: Session = Depends(get_db)
):
    """Create new course content record manually."""
    db_content = CourseContent(**content.dict())
    db.add(db_content)
    db.commit()
    db.refresh(db_content)
    return db_content


@router.get("/", response_model=List[CourseContentResponse])
async def get_all_course_content(
    skip: int = 0,
    limit: int = 100,
    db: Session = Depends(get_db)
):
    """Get all course content records."""
    content = db.query(CourseContent).offset(skip).limit(limit).all()
    return content


@router.get("/{content_id}", response_model=CourseContentResponse)
async def get_course_content(content_id: int, db: Session = Depends(get_db)):
    """Get specific course content by ID."""
    content = db.query(CourseContent).filter(CourseContent.id == content_id).first()
    if not content:
        raise HTTPException(status_code=404, detail="Course content not found")
    return content


@router.put("/{content_id}", response_model=CourseContentResponse)
async def update_course_content(
    content_id: int,
    content_update: CourseContentUpdate,
    db: Session = Depends(get_db)
):
    """Update course content."""
    db_content = db.query(CourseContent).filter(CourseContent.id == content_id).first()
    if not db_content:
        raise HTTPException(status_code=404, detail="Course content not found")
    
    update_data = content_update.dict(exclude_unset=True)
    for field, value in update_data.items():
        setattr(db_content, field, value)
    
    db.commit()
    db.refresh(db_content)
    return db_content


@router.delete("/{content_id}")
async def delete_course_content(content_id: int, db: Session = Depends(get_db)):
    """Delete course content."""
    db_content = db.query(CourseContent).filter(CourseContent.id == content_id).first()
    if not db_content:
        raise HTTPException(status_code=404, detail="Course content not found")
    
    db.delete(db_content)
    db.commit()
    return {"message": "Course content deleted successfully"}


@router.post("/{content_id}/generate-summary")
async def generate_content_summary(
    content_id: int,
    custom_prompt: Optional[str] = None,
    db: Session = Depends(get_db)
):
    """Generate AI summary for course content."""
    
    # Get the course content
    db_content = db.query(CourseContent).filter(CourseContent.id == content_id).first()
    if not db_content:
        raise HTTPException(status_code=404, detail="Course content not found")
    
    try:
        # Generate summary using Gemini AI
        summary = await gemini_service.generate_content_summary(
            content=db_content.content,
            custom_prompt=custom_prompt or ""
        )
        
        # Update the content record with the summary
        db_content.summary = summary
        db.commit()
        db.refresh(db_content)
        
        return {
            "content_id": content_id,
            "summary": summary,
            "message": "Summary generated successfully"
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating summary: {str(e)}") 